"""
app/ingestion/docx_parser.py

DOCX text extraction via python-docx. Extracts paragraphs and table
content. Does not extract embedded images — out of scope for Phase 1
text-only RAG.
"""

import asyncio
import io

import structlog
from docx import Document as DocxDocument

from app.core.exceptions import ParserError
from app.ingestion.base_parser import BaseParser, ParsedDocument

logger = structlog.get_logger(__name__)


class DOCXParser(BaseParser):
    async def parse(self, file_bytes: bytes) -> ParsedDocument:
        try:
            return await asyncio.to_thread(self._parse_sync, file_bytes)
        except ParserError:
            raise
        except Exception as e:
            logger.error("docx_parse_failed", error=str(e))
            raise ParserError("Could not extract text from DOCX file. File may be corrupted.") from e

    def _parse_sync(self, file_bytes: bytes) -> ParsedDocument:
        doc = DocxDocument(io.BytesIO(file_bytes))

        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)

        full_text = "\n".join(parts)
        if not full_text.strip():
            raise ParserError("DOCX file contains no extractable text.")

        return ParsedDocument(
            text=full_text,
            page_count=1,  # DOCX has no native page concept until rendered
            pages=[full_text],
            metadata={"parser": "python-docx"},
        )